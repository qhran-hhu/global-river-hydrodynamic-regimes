# -*- coding: utf-8 -*-
"""汇编补充材料 S1-S5（docx + md）。

S1 构造无关稳健性矩阵（特征剔除 × 模态检验）
S2 短记录保真度（D2 两臂）
S3 噪声衰减校正（D3 两臂）
S4 端元三亚组画像（C2）
S5 QC 阈值敏感性（C1）
"""
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).resolve().parent
RS = BASE.parent / "code" / "output" / "regime_space"
DOCX = BASE / "supplementary_S1-S5.docx"
MD = BASE / "supplementary_S1-S5.md"

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
md = []


def h(text, level=1):
    doc.add_heading(text, level=level)
    md.append("#" * level + " " + text + "\n")


def p(text):
    doc.add_paragraph(text)
    md.append(text + "\n")


def table(df, caption, max_rows=40):
    doc.add_paragraph(caption).runs[0].bold = True
    md.append("\n**" + caption + "**\n")
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for j, c in enumerate(df.columns):
        t.rows[0].cells[j].text = str(c)
    md.append("| " + " | ".join(str(c) for c in df.columns) + " |")
    md.append("|" + "---|" * len(df.columns))
    for _, r in df.head(max_rows).iterrows():
        cells = t.add_row().cells
        md.append("| " + " | ".join(str(v) for v in r.values) + " |")
        for j, v in enumerate(r.values):
            cells[j].text = str(v)
    doc.add_paragraph("")
    md.append("")


# ================= 封面 =================
ti = doc.add_heading("Supplementary Information", level=0)
ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
p("The global hydrodynamic-regime continuum: supplementary robustness notes "
  "S1-S5")
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p("Qihua Ran — Hohai University")
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
md.append("# Supplementary Information S1-S5\n\nQihua Ran — Hohai University\n")

# ================= S1 =================
h("Note S1. Construction-independent robustness of the continuum", 1)
p("The three modality tests (Hartigan dip on PC1/PC2 with bootstrap p-values; "
  "Gaussian-mixture BIC improvement at k = 2 over k = 1 on a 50,000-reach "
  "subsample, fitted to convergence (ten initialisations) and calibrated "
  "against unimodal references with identical mean and covariance (Gaussian "
  "and kurtosis-matched heavy-tailed t); "
  "scale-space mode persistence) were repeated on the full feature "
  "set and under every feature ablation: without the two f5 phase dimensions "
  "(nof5), without f4 (nof4), on the slope-free set {f2, f3, f5} (noslope), "
  "and on the same set restricted to reaches with a measured annual "
  "harmonic (f5 not gated; noslope_nogate). In every variant the dip tests "
  "remain non-significant, the k = 2 BIC improvement stays at 1.0-2.2% "
  "(approximately 0% for the unimodal references), and a single persistent "
  "mode is reached by bandwidths of 2-3 histogram cells.")
s1 = pd.read_csv(RS / "特征剔除稳健性矩阵.csv")
s1.columns = ["variant", "test", "stat", "detail"]
table(s1, "Table S1. Feature-ablation robustness matrix (patched f5; "
          "dip entries give subsample-median bootstrap p; gmm_real entries "
          "give BIC improvement at k = 2 and k = 6 in %).")
s1b = pd.read_csv(RS / "noslope归因_r2.csv")
table(s1b.round(3), "Table S1b. Driver attribution in the slope-free regime "
                    "space (held-out R2; slope fully exogenous).")
rob = pd.read_csv(RS / "构造无关稳健性.csv")
rob.columns = ["test", "dip_stat", "p_median"][:len(rob.columns)] if len(rob.columns) == 3 else rob.columns
table(rob.round(3), "Table S1c. Additional modality diagnostics "
                    "(marginal dip tests, robust-rescaled PCs, threshold "
                    "sweep of the seasonal-dominance fraction, PC-space "
                    "HDBSCAN).")

# ================= S2 =================
h("Note S2. Snapshot fidelity: does a 3.4-year record suffice?", 1)
p("Arm A — long GSIM records. On 16,546 stations with at least 20 years of "
  "monthly discharge, five random 41-month (3.4-year) windows per station "
  "(79,443 windows) reproduce the full-record regime features at rank "
  "correlations of 0.86 (relative range), 0.88 (IQR of log Q), 0.86 (event "
  "response), 0.97 (annual-harmonic amplitude) and 0.76 (harmonic R2); "
  "seasonal phase is recovered to a median absolute error of 0.20 months "
  "(98.7% within one month, 99.9% within two) where both harmonics are "
  "significant (n = 26,245).")
p("Arm B — SWOT split halves. Each reach's 3.4-year record was split at its "
  "median date into independent 1.7-year halves (130,989 reaches with at "
  "least 12 valid passes in both halves). Half-to-half rank correlations: "
  "f2 0.534 (the most event-sensitive feature), f3 0.898, f4 0.653, "
  "amplitude 0.787, R2 0.698; seasonal phase "
  "agrees to a median 0.36 months (89.4% within one month, 98.7% within two; "
  "n = 32,845 with both halves at R2 >= 0.3). Half-record features reproduce "
  "the full-record features at rank 0.79-0.96 (f2 0.79, f3 0.94-0.96, f4 "
  "0.85, amplitude 0.91, R2 0.88). Regime features are thus far more stable "
  "within the observation window than the cross-variable (discharge versus "
  "hydrodynamic) coupling reported in the main text: record length is not a "
  "limiting factor for the conclusions.")
p("Full statistics: D2_快照保真度.txt in the project pipeline "
  "(window-level detail: D2_gsim_windows.csv, 79,443 rows).")

# ================= S3 =================
h("Note S3. Measurement-noise attenuation", 1)
p("Arm 1 — noise injection. Simulated SWOT noise was added to the full time "
  "series of a stratified 2,000-reach sample (continent x variability "
  "quartiles): additive Gaussian WSE error at sigma = 0.1 m (the reach-scale "
  "level measured in Cal/Val), 0.2 m and 0.3 m (pessimistic), with "
  "multiplicative lognormal width error of 10-20%; five replicates per "
  "scenario. Feature ranks are essentially unchanged (f2: Spearman rho = "
  "0.98, 0.95 and 0.92 at sigma = 0.1, 0.2 and 0.3 m; f4 the most sensitive, "
  "0.88 at the pessimistic level); the annual-harmonic R2 "
  "drops by at most 0.002; phase stays within one month for 100% of "
  "seasonally dominated reaches. Crucially, dip tests on the noise-added "
  "regime space remain non-significant at every noise level (p >= 0.94): "
  "unimodality is not produced by noise smoothing.")
s3 = pd.read_csv(RS / "D3_噪声衰减.csv")
agg = s3.groupby(["sigma_wse", "sigma_width"]).median(numeric_only=True)
agg = agg.drop(columns=["rep", "n"], errors="ignore").round(3).reset_index()
table(agg, "Table S3. Noise-injection summary (median of five replicates; "
           "rho = rank correlation against the clean features; bias in %).")
p("Arm 2 — disattenuation. Split-half reliabilities (Spearman-Brown "
  "extrapolated to the full record): f2 0.697, f3 0.946, f4 0.790, "
  "amplitude 0.881, R2 0.822. Correcting the hydrodynamic-discharge coupling "
  "for attenuation on both sides raises rho(f4, IQR log Q) from 0.128 to "
  "0.154 and rho(f2, q2) from 0.095 to 0.122 — the weak coupling in Section "
  "3.5 is not a noise artifact.")

# ================= S4 =================
h("Note S4. Anatomy of the weak-annual-cycle end", 1)
p("Reaches below the seasonal-dominance threshold (f5 R2 < 0.3; 71.5% of the "
  "quality-controlled sample) partition into three named subgroups plus a "
  "residual transitional group (mutually exclusive, in priority order "
  "cold -> dam -> arid): cold/ice-affected reaches (mean annual temperature "
  "below 0 degC; median latitude 63, largest width event response, f3 = "
  "3.4), dam-adjacent reaches (within the dam neighbourhood by construction; "
  "depth-relative water-level range unremarkable, f2 = 1.91 versus 1.86 in "
  "the transitional group — consistent with the minimal median dam effect of "
  "Section 3.4), and arid event-driven reaches (annual precipitation below "
  "500 mm; lowest variability overall, f2 = 1.80). The strong-seasonal end "
  "(R2 >= 0.6; "
  "8.9% of reaches, tropical-dominated) is shown for contrast. The residual "
  "'transitional' group (40.8%) sits between the end-members on every "
  "gradient — it is the continuum's interior, not a hidden class.")
s4 = pd.read_csv(RS / "C2_端元三亚组画像.csv")
s4.columns = ["subgroup", "n", "share_%", "f2_med", "f3_med", "f4_med",
              "f5R2_med", "lat_med", "bio1_meanT", "bio12_annualP",
              "bio15_Pseason", "ice_frac_med", "dam_adj_%",
              "top_continents"]
s4["subgroup"] = s4["subgroup"].map({
    "其他过渡": "transitional (residual)",
    "寒冷/冰冻影响": "cold/ice-affected",
    "干旱事件型": "arid event-driven",
    "坝调控": "dam-regulated",
    "强季节端（对照）": "strong-seasonal end (ref.)"})
table(s4, "Table S4. End-member subgroup profiles (medians unless noted).")

# ================= S5 =================
h("Note S5. Sensitivity to quality-control thresholds", 1)
p("The full modality battery was repeated while varying each quality-control "
  "threshold (climatological ice fraction, dark-pixel fraction, minimum "
  "observation count), including a no-QC variant. The continuum is "
  "insensitive to every choice: dip tests remain non-significant "
  "(p >= 0.85) and the k = 2 BIC improvement stays below 3.7% in all "
  "variants, including the 40,000-reach subsample with the strictest "
  "observation-count filter.")
s5 = pd.read_csv(RS / "C1_QC阈值敏感性.csv")
s5 = s5.round({"dip_pc1_p": 3, "dip_pc2_p": 3, "gmm_k2_impr_pct": 2})
table(s5, "Table S5. QC-threshold sensitivity of the modality tests.")

doc.save(DOCX)
MD.write_text("\n".join(md), encoding="utf-8")
print("saved:", DOCX)
print("saved:", MD)
